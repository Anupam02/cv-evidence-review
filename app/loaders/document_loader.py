"""
Step 1: Document Loading
-------------------------
Same problem as rag-service's loader: turn a file (PDF/docx/txt) into
clean plain text. Reused here almost verbatim because it's genuinely
the same problem - a CV is just another document.

One CV-specific addition: we keep a simple page count / structure hint
in metadata, because CV layout quality (single column vs. multi-column,
tables for skills, etc.) directly affects how reliably we can later
validate that an "evidence quote" the LLM extracted actually appears
in the source text. A CV that extracts as garbled/reordered text due
to a complex layout is a real limitation worth being able to flag.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class Document:
    text: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


def load_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return Document(text=text, source=path.name)


def load_pdf(path: Path) -> Document:
    import pdfplumber

    pages_text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")

    full_text = "\n\n".join(pages_text)
    return Document(
        text=full_text,
        source=path.name,
        metadata={"num_pages": len(pages_text)},
    )


def load_docx(path: Path) -> Document:
    import docx

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return Document(text=full_text, source=path.name)


LOADERS = {
    ".txt": load_txt,
    ".md": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(path: str | Path) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in LOADERS:
        raise ValueError(
            f"Unsupported file type: {suffix}. Supported: {list(LOADERS.keys())}"
        )
    return LOADERS[suffix](path)


if __name__ == "__main__":
    import sys
    doc = load_document(sys.argv[1])
    print(f"Loaded {doc.source}: {len(doc.text)} chars")
    print(doc.metadata)
    print("---first 400 chars---")
    print(doc.text[:400])
